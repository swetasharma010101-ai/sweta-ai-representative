"""
Ingestion script — run once to build the ChromaDB vector store.
Uses FREE local HuggingFace embeddings (no API key needed).

    python -m rag.ingest
"""

import os
import sys
from pathlib import Path
from typing import List
from dotenv import load_dotenv
load_dotenv()

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

try:
    from github import Github
except ImportError:
    Github = None

CHROMA_DIR  = os.getenv("CHROMA_PERSIST_DIR", "./chroma_db")
EMBED_MODEL = os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2")
COLLECTION  = "sweta_persona"
DATA_DIR    = Path(__file__).parent.parent.parent / "data"

splitter = RecursiveCharacterTextSplitter(
    chunk_size=600,
    chunk_overlap=80,
    separators=["\n\n", "\n", ". ", " ", ""],
)


def load_resume() -> List[Document]:
    docs = []
    for ext, loader in [("pdf", _load_pdf), ("docx", _load_docx), ("md", _load_text), ("txt", _load_text)]:
        path = DATA_DIR / f"resume.{ext}"
        if path.exists():
            text = loader(path)
            docs.append(Document(
                page_content=text,
                metadata={"source": f"resume.{ext}", "type": "resume"}
            ))
            print(f"  ✓ Loaded resume: {path.name} ({len(text)} chars)")
            break
    if not docs:
        print("  ⚠️  No resume found in data/. Add resume.pdf, resume.docx, or resume.md")
    return docs


def _load_pdf(path):
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def _load_docx(path):
    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs)

def _load_text(path):
    return path.read_text(encoding="utf-8")


def load_extra_docs() -> List[Document]:
    docs = []
    for p in DATA_DIR.glob("*.md"):
        if p.stem == "resume":
            continue
        text = p.read_text(encoding="utf-8")
        docs.append(Document(page_content=text, metadata={"source": p.name, "type": "extra"}))
        print(f"  ✓ Loaded extra doc: {p.name}")
    return docs


def load_github_repos() -> List[Document]:
    token    = os.getenv("GITHUB_TOKEN")
    username = os.getenv("GITHUB_USERNAME", "swetasharma")
    if not token or not Github:
        print("  ⚠️  GITHUB_TOKEN not set — skipping GitHub ingestion")
        return []

    g     = Github(token)
    user  = g.get_user(username)
    docs  = []
    repos = [r for r in list(user.get_repos(type="owner", sort="updated"))[:15] if not r.fork]

    for repo in repos:
        print(f"  ✓ Ingesting: {repo.name}")
        try:
            readme = repo.get_readme()
            content = readme.decoded_content.decode("utf-8")
            docs.append(Document(
                page_content=f"# {repo.name} README\n\n{content}",
                metadata={"source": f"github/{repo.name}/README", "type": "github_readme", "repo": repo.name},
            ))
        except Exception:
            pass

        meta = (
            f"Repo: {repo.name}\nDescription: {repo.description or 'N/A'}\n"
            f"Language: {repo.language or 'N/A'}\nStars: {repo.stargazers_count}\n"
            f"URL: {repo.html_url}\n"
        )
        docs.append(Document(
            page_content=meta,
            metadata={"source": f"github/{repo.name}/meta", "type": "github_meta"},
        ))

    return docs


def main():
    print("\n🔍 Starting ingestion...\n")
    raw_docs: List[Document] = []

    print("📄 Loading resume...")
    raw_docs += load_resume()

    print("📁 Loading extra docs...")
    raw_docs += load_extra_docs()

    print("🐙 Loading GitHub repos...")
    raw_docs += load_github_repos()

    if not raw_docs:
        print("\n❌ No documents found. Add data/resume.md at minimum.")
        sys.exit(1)

    print(f"\n✂️  Splitting {len(raw_docs)} documents...")
    chunks = splitter.split_documents(raw_docs)
    print(f"   → {len(chunks)} chunks")

    print(f"\n🤖 Loading local embedding model (first run downloads ~90MB)...")
    embeddings = HuggingFaceEmbeddings(
        model_name=EMBED_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    print(f"\n💾 Storing in ChromaDB at {CHROMA_DIR}...")
    Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=COLLECTION,
        persist_directory=CHROMA_DIR,
    )
    print(f"\n✅ Done! {len(chunks)} chunks stored.\n")


if __name__ == "__main__":
    main()
